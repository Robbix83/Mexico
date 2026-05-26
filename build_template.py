"""
build_template.py — Builds templates/cctv.docx from scratch using python-docx.

The template uses docxtemplater placeholder syntax:
  {field_name}              — simple text substitution
  {#loop_name}...{/loop_name} — array iteration
  {%image_tag}              — image embedding (image-module-free)

This is a STARTER template — Hebrew RTL is enabled per paragraph, fonts default
to Arial. The user can refine fonts/colors/cover by editing the resulting docx
in Word (placeholders survive normal Word editing as long as they stay in one
text-run).
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_rtl(paragraph):
    """Enable RTL on a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)


def set_run_rtl(run):
    """Add <w:rtl/> to a run's rPr.
    Paragraph-level <w:bidi/> is respected by Word but largely ignored by
    Google Docs. Run-level <w:rtl/> is respected by both — this is the
    reliable cross-app way to force RTL on each text run."""
    rPr = run._r.get_or_add_rPr()
    for e in rPr.findall(qn('w:rtl')):
        rPr.remove(e)
    rtl = OxmlElement('w:rtl')
    rPr.append(rtl)


def set_style_run_rtl(style):
    """Add <w:rtl/> to a style's rPr so all runs using this style
    (including ones docxtemplater creates at render time) inherit RTL."""
    rPr = style.element.get_or_add_rPr()
    for e in rPr.findall(qn('w:rtl')):
        rPr.remove(e)
    rtl = OxmlElement('w:rtl')
    rPr.append(rtl)


def set_table_rtl(table):
    """Enable RTL on a table."""
    tblPr = table._tbl.tblPr
    bidiVisual = OxmlElement('w:bidiVisual')
    bidiVisual.set(qn('w:val'), '1')
    tblPr.append(bidiVisual)


def set_section_rtl(section):
    """Enable RTL at the section level so it cascades to all content,
    including paragraphs that docxtemplater creates dynamically inside
    {#loops} or where it inserts images."""
    sectPr = section._sectPr
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    # Insert as the first child so Word reads it before page setup
    sectPr.insert(0, bidi)


def set_style_rtl(style):
    """Enable RTL on a style's paragraph-properties so any paragraph
    using this style (including ones created at render time) inherits
    bidi without us having to set it per-paragraph."""
    pPr = style.element.get_or_add_pPr()
    # Avoid duplicating <w:bidi/> if already present
    existing = pPr.findall(qn('w:bidi'))
    for e in existing:
        pPr.remove(e)
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    # Also pin right alignment on the style itself (jc=right is RTL's "natural left")
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc')
        pPr.append(jc)
    jc.set(qn('w:val'), 'right')


def shade_cell(cell, color_hex):
    """Apply background shading to a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def add_heading_rtl(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    set_rtl(h)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in h.runs:
        set_run_rtl(run)
    return h


def add_para_rtl(doc, text='', bold=False, size=None, align=WD_ALIGN_PARAGRAPH.RIGHT):
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_rtl(run)
        if bold:
            run.bold = True
        if size:
            run.font.size = Pt(size)
        run.font.name = 'Arial'
    return p


def add_bullet_rtl(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if text:
        run = p.add_run(text)
        set_run_rtl(run)
        run.font.name = 'Arial'
    return p


def set_document_hebrew(doc):
    """Tell Word this document's primary language is Hebrew (he-IL).
    Without this, the default <w:lang ... w:bidi="ar-SA"/> in rPrDefault makes
    Word treat the document as English-with-Arabic-bidi — Hebrew text still
    appears but Word's reading-order heuristics misalign the page.
    Fix: rewrite both rPrDefault/<w:lang> in styles.xml AND themeFontLang in
    settings.xml at the python-docx XML level before saving.
    """
    # rPrDefault > w:rPr > w:lang
    docDefaults = doc.styles.element.find(qn('w:docDefaults'))
    if docDefaults is not None:
        rPrDef = docDefaults.find(qn('w:rPrDefault'))
        if rPrDef is not None:
            rPr = rPrDef.find(qn('w:rPr'))
            if rPr is not None:
                lang = rPr.find(qn('w:lang'))
                if lang is None:
                    lang = OxmlElement('w:lang')
                    rPr.append(lang)
                lang.set(qn('w:val'),     'he-IL')
                lang.set(qn('w:eastAsia'),'he-IL')
                lang.set(qn('w:bidi'),    'he-IL')
            # Also stamp <w:rtl/> on rPrDefault so every run that inherits
            # the document default is RTL — Google Docs reads this.
            if rPr is not None:
                for e in rPr.findall(qn('w:rtl')):
                    rPr.remove(e)
                rPr.append(OxmlElement('w:rtl'))
    # settings.xml > w:themeFontLang
    settings = doc.settings.element
    tfl = settings.find(qn('w:themeFontLang'))
    if tfl is None:
        tfl = OxmlElement('w:themeFontLang')
        settings.append(tfl)
    tfl.set(qn('w:val'),     'he-IL')
    tfl.set(qn('w:eastAsia'),'he-IL')
    tfl.set(qn('w:bidi'),    'he-IL')


def build():
    doc = Document()

    # ── RTL EVERYWHERE ──
    # 1. Document language = Hebrew (he-IL). CRITICAL: without this Word
    #    interprets <w:bidi/> via Arabic conventions (the default bidi lang
    #    is ar-SA) and the page-level reading order is wrong.
    set_document_hebrew(doc)

    # 2. Section-level bidi: cascades to anything Word can't otherwise
    #    figure out (e.g. paragraphs docxtemplater fabricates inside loops
    #    or where it inlines images). Without this, any paragraph that
    #    inherits from "Normal" w/o its own <w:bidi/> renders LTR.
    for sec in doc.sections:
        set_section_rtl(sec)

    # 3. Style-level bidi (paragraph) + rtl (run) on every style we use.
    #    set_style_rtl  → <w:bidi/> + jc=right on pPr  (Word paragraph RTL)
    #    set_style_run_rtl → <w:rtl/> on rPr (Google Docs run RTL — critical)
    for style_name in ('Normal', 'Heading 1', 'Heading 2', 'Heading 3', 'List Bullet',
                       'Default Paragraph Font'):
        try:
            set_style_rtl(doc.styles[style_name])
            set_style_run_rtl(doc.styles[style_name])
        except KeyError:
            pass  # style not present in this build

    # Default font: Arial 11
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # ═════════════════ COVER ═════════════════
    p = add_para_rtl(doc, '', align=WD_ALIGN_PARAGRAPH.CENTER)
    p.add_run('\n\n\n').font.size = Pt(20)
    # Centered cover title — Hebrew only, no Latin chars to avoid BIDI reorder
    p = add_para_rtl(doc, 'תיק תיעוד', align=WD_ALIGN_PARAGRAPH.CENTER)
    p.runs[0].font.size = Pt(36)
    p.runs[0].bold = True
    p.runs[0].font.color.rgb = RGBColor(0x1a, 0x2a, 0x4a)
    p = add_para_rtl(doc, 'AS-MADE', align=WD_ALIGN_PARAGRAPH.CENTER)
    p.runs[0].font.size = Pt(20)
    p.runs[0].font.color.rgb = RGBColor(0x60, 0x70, 0x80)

    p = add_para_rtl(doc, '{site_name}', align=WD_ALIGN_PARAGRAPH.CENTER)
    p.runs[0].font.size = Pt(28)
    p.runs[0].bold = True

    p = add_para_rtl(doc, '{site_address}', align=WD_ALIGN_PARAGRAPH.CENTER)
    p.runs[0].font.size = Pt(16)

    p = add_para_rtl(doc, '{site_subtitle}', align=WD_ALIGN_PARAGRAPH.CENTER)
    p.runs[0].font.size = Pt(12)
    p.runs[0].italic = True

    # ── Cover-page project info table ──
    add_para_rtl(doc, '\n')
    info_table = doc.add_table(rows=7, cols=4)
    info_table.style = 'Light Grid Accent 1'
    # bidiVisual is NOT used here: Google Docs ignores <w:bidiVisual/> and renders
    # columns in physical left→right order regardless.  Instead we physically reverse
    # the column order so the role-label (rightmost in Hebrew) sits in physical col 3
    # (rightmost in both Word and Google Docs without any bidi table flag).
    # Reading order for a Hebrew reader (right → left):
    #   col3 (role/label) | col2 (function) | col1 (org) | col0 (person)

    rows_data = [
        ('{contractor_person}', '{contractor_org}', 'גורם מבצע',      'מגיש'),
        ('{pm_person}',         '{pm_org}',          'גורם מוביל',     'מגיש'),
        ('{customer_person}',   '{customer_org}',    'לקוח סופי',      'מקבל'),
        ('{consultant_person}', '{consultant_org}',  'יועץ טכנולוגי',  'בודק'),
        ('{engineer_person}',   '{engineer_org}',    'מהנדס טכנולוגי', 'מאשר'),
        ('',                    '{submit_date}',     'תאריך הגשה',     'תאריך'),
        ('',                    '{update_date}',     'תאריך עדכון',    'תאריך'),
    ]
    for i, row_vals in enumerate(rows_data):
        cells = info_table.rows[i].cells
        for cell, txt in zip(cells, row_vals):
            cell.text = ''
            p = cell.paragraphs[0]
            set_rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            run = p.add_run(txt)
            set_run_rtl(run)
            run.font.name = 'Arial'
            run.font.size = Pt(11)
        shade_cell(cells[3], 'E8EEF7')  # role label — physically rightmost col
        shade_cell(cells[2], 'F4F6FA')  # function  — second from right

    doc.add_page_break()

    # ═════════════════ 1. הקדמה ═════════════════
    # Drop numeric prefixes from headings — Latin digits ("1.", "2."...) inside
    # an RTL paragraph get reordered by Unicode BIDI to the visual left, which
    # made the heading look reversed in Word ("הקדמה 1." instead of "1. הקדמה").
    # Plain Hebrew titles render cleanly.
    add_heading_rtl(doc, 'הקדמה', level=1)
    add_para_rtl(doc, '{intro_text}')
    add_para_rtl(doc, '')

    # ═════════════════ 2. תכולת הפרויקט ═════════════════
    add_heading_rtl(doc, 'תכולת הפרויקט', level=1)
    # Loop over scope_items
    p = add_bullet_rtl(doc, '{#scope_items}{text}{/scope_items}')
    add_para_rtl(doc, '')

    # ═════════════════ 3. ארכיטקטורת רשת ═════════════════
    add_heading_rtl(doc, 'ארכיטקטורת רשת התקשורת', level=1)
    add_para_rtl(doc, '{network_arch_text}')
    add_para_rtl(doc, '')
    # Network diagram image
    add_para_rtl(doc, 'דיאגרמת רשת:', bold=True)
    p = add_para_rtl(doc, '{%network_diagram_img}', align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para_rtl(doc, '')

    # ═════════════════ 4. שרטוט AS-MADE ═════════════════
    add_heading_rtl(doc, 'שרטוט AS-MADE', level=1)
    add_para_rtl(doc, 'תכנית האתר:', bold=True)
    add_para_rtl(doc, '{%site_plan_img}', align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para_rtl(doc, '')

    add_para_rtl(doc, 'תמונות מהשטח:', bold=True)
    # Photos loop
    add_para_rtl(doc, '{#photos}', align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para_rtl(doc, '{%img}', align=WD_ALIGN_PARAGRAPH.CENTER)
    p = add_para_rtl(doc, '{caption}', align=WD_ALIGN_PARAGRAPH.CENTER)
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(10)
    add_para_rtl(doc, '{/photos}')

    doc.add_page_break()

    # ═════════════════ 5. רשימת ציוד וכתובות IP ═════════════════
    add_heading_rtl(doc, 'רשימת ציוד וכתובות IP', level=1)

    # ── 5.1 Cameras ──
    add_heading_rtl(doc, 'רשימת מצלמות', level=2)
    cam_table = doc.add_table(rows=2, cols=7)
    cam_table.style = 'Light Grid Accent 1'
    # Columns physically reversed so they display in correct Hebrew R→L reading order
    # without needing <w:bidiVisual/> (which Google Docs ignores).
    # Physical L→R: מיקום | IP | פורט | סוג | שם | ארון | מס"ד
    # Hebrew reader sees R→L: מס"ד | ארון | שם | סוג | פורט | IP | מיקום ✓
    cam_headers = ['מיקום בתכנית', 'כתובת IP', 'פורט בפנל', 'סוג מצלמה', 'שם המצלמה', 'ארון תקשורת', 'מס"ד']
    for cell, h in zip(cam_table.rows[0].cells, cam_headers):
        cell.text = ''
        p = cell.paragraphs[0]
        set_rtl(p); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.bold = True; r.font.name = 'Arial'; r.font.size = Pt(10)
        set_run_rtl(r)
        shade_cell(cell, '1A2A4A')
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Data row — loop fields match the reversed header order
    cam_data_cells = cam_table.rows[1].cells
    cam_loop_fields = [
        '{#cameras}{location}', '{ip}', '{port}', '{model}', '{name}', '{cabinet}', '{idx}{/cameras}',
    ]
    for cell, txt in zip(cam_data_cells, cam_loop_fields):
        cell.text = ''
        p = cell.paragraphs[0]
        set_rtl(p); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt); r.font.name = 'Arial'; r.font.size = Pt(10)
        set_run_rtl(r)
    add_para_rtl(doc, '')

    # ── 5.2 Backhauls ──
    add_heading_rtl(doc, 'רשימת עורקים', level=2)
    bh_table = doc.add_table(rows=2, cols=6)
    bh_table.style = 'Light Grid Accent 1'
    # Columns physically reversed (L→R: IP | מיקום | יצרן | מק"ט | סוג עורק | מס"ד)
    bh_headers = ['כתובת IP', 'מיקום', 'יצרן', 'מק"ט', 'סוג עורק', 'מס"ד']
    for cell, h in zip(bh_table.rows[0].cells, bh_headers):
        cell.text = ''
        p = cell.paragraphs[0]
        set_rtl(p); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.bold = True; r.font.name = 'Arial'; r.font.size = Pt(10)
        set_run_rtl(r)
        shade_cell(cell, '1A2A4A')
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    bh_data_cells = bh_table.rows[1].cells
    bh_loop_fields = [
        '{#backhauls}{ip}', '{location}', '{vendor}', '{mpn}', '{type}', '{idx}{/backhauls}',
    ]
    for cell, txt in zip(bh_data_cells, bh_loop_fields):
        cell.text = ''
        p = cell.paragraphs[0]
        set_rtl(p); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt); r.font.name = 'Arial'; r.font.size = Pt(10)
        set_run_rtl(r)
    add_para_rtl(doc, '')

    # ── 5.3 Switches ──
    add_heading_rtl(doc, 'רשימת מתגים', level=2)
    sw_table = doc.add_table(rows=2, cols=5)
    sw_table.style = 'Light Grid Accent 1'
    # Columns physically reversed (L→R: IP | יצרן | מק"ט | מתגים | מס"ד)
    sw_headers = ['כתובת IP', 'יצרן', 'מק"ט', 'מתגים', 'מס"ד']
    for cell, h in zip(sw_table.rows[0].cells, sw_headers):
        cell.text = ''
        p = cell.paragraphs[0]
        set_rtl(p); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.bold = True; r.font.name = 'Arial'; r.font.size = Pt(10)
        set_run_rtl(r)
        shade_cell(cell, '1A2A4A')
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    sw_data_cells = sw_table.rows[1].cells
    sw_loop_fields = [
        '{#switches}{ip}', '{vendor}', '{mpn}', '{name}', '{idx}{/switches}',
    ]
    for cell, txt in zip(sw_data_cells, sw_loop_fields):
        cell.text = ''
        p = cell.paragraphs[0]
        set_rtl(p); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt); r.font.name = 'Arial'; r.font.size = Pt(10)
        set_run_rtl(r)

    # ═════════════════ 6. נספח — דפי מוצר ═════════════════
    # Each datasheet PDF is rendered to PNG pages by the server at export time
    # and inlined here. The outer loop iterates datasheets; the inner loop
    # iterates pages within a datasheet. paragraphLoop:true (set in docpack.js)
    # makes loops that wrap whole paragraphs repeat each contained paragraph.
    doc.add_page_break()
    add_heading_rtl(doc, 'נספח — דפי מוצר', level=1)
    p = add_para_rtl(doc,
        'דפי המוצר של כל הרכיבים בפרויקט. כל דאטה-שיט מופיע בעמוד נפרד.')
    p.runs[0].font.size = Pt(10.5)
    add_para_rtl(doc, '')

    # Outer loop: one entry per datasheet; inner loop renders pages as images
    add_para_rtl(doc, '{#datasheets}')
    # Inner loop: each page is one full-width image — no text labels
    add_para_rtl(doc, '{#pages}')
    add_para_rtl(doc, '{%page_img}', align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para_rtl(doc, '{/pages}')
    add_para_rtl(doc, '')
    add_para_rtl(doc, '{/datasheets}')

    # ── Footer ──
    add_para_rtl(doc, '')
    add_para_rtl(doc, '')
    p = add_para_rtl(doc, 'תיק תיעוד זה הופק אוטומטית ממערכת אפקון.', align=WD_ALIGN_PARAGRAPH.CENTER)
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(0x60, 0x70, 0x80)

    # Save
    import os
    out = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'templates', 'cctv.docx')
    doc.save(out)
    print(f'Wrote {out}')


if __name__ == '__main__':
    build()
